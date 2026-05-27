#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera Trabajo_Parcial_Cap1_Cap2_v2.docx
Aplica cambios CP1–CP6 sobre el contenido del v1
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name='Times New Roman', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force font for complex scripts (needed for non-latin)
    r = run._r
    rpr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    existing = rpr.find(qn('w:rFonts'))
    if existing is not None:
        rpr.remove(existing)
    rpr.insert(0, rFonts)

def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             bold=False, size=12, space_before=0, space_after=200,
             first_indent=False, color=None, heading_level=None):
    """Add a paragraph with standard formatting."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Cm(1.25)
    pf.line_spacing = Pt(21.6)  # 1.5 × 14.4pt (approx 1.5 for 12pt)

    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p

def add_heading(doc, text, level=1):
    """Section heading: bold, centered for chapter titles, left for subsections."""
    if level == 0:  # Chapter title
        return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER,
                       bold=True, size=13, space_before=18, space_after=12)
    elif level == 1:
        return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       bold=True, size=12, space_before=14, space_after=8)
    else:
        return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       bold=True, size=12, space_before=8, space_after=6)

def add_body(doc, text, first_indent=True):
    """Standard body paragraph."""
    return add_para(doc, text, first_indent=first_indent,
                   space_before=0, space_after=10)

def add_objective(doc, text):
    """Objective paragraph (bold label + normal text)."""
    parts = text.split(' — ', 1) if ' — ' in text else [text, '']
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = Pt(21.6)
    if parts[1]:
        r1 = p.add_run(parts[0] + ' — ')
        set_font(r1, bold=True)
        r2 = p.add_run(parts[1])
        set_font(r2, bold=False)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

def add_header_footer(doc):
    section = doc.sections[0]
    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run('MIA-303 – Trabajo de Investigación I | Carlos Pérez Pérez')
    hr.font.name = 'Times New Roman'
    hr.font.size = Pt(10)
    hr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Footer with page number
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('Página ')
    fr.font.name = 'Times New Roman'
    fr.font.size = Pt(10)
    # page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_fld = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20')
    rpr.append(sz)
    run_fld.append(rpr)
    run_fld.append(fldChar1)
    run_fld.append(instrText)
    run_fld.append(fldChar2)
    fp._p.append(run_fld)
    fr2 = fp.add_run(' de ')
    fr2.font.name = 'Times New Roman'
    fr2.font.size = Pt(10)
    # total pages field
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = ' NUMPAGES '
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run_fld2 = OxmlElement('w:r')
    rpr2 = OxmlElement('w:rPr')
    sz2 = OxmlElement('w:sz'); sz2.set(qn('w:val'), '20')
    rpr2.append(sz2)
    run_fld2.append(rpr2)
    run_fld2.append(fldChar3)
    run_fld2.append(instrText2)
    run_fld2.append(fldChar4)
    fp._p.append(run_fld2)

# ─── Document content ────────────────────────────────────────────────────────

doc = Document()
set_page_margins(doc)
# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

add_header_footer(doc)

# ── PORTADA ──────────────────────────────────────────────────────────────────
add_para(doc, 'UNIVERSIDAD NACIONAL DE INGENIERÍA',
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14,
         space_before=0, space_after=6)
add_para(doc, 'FACULTAD DE INGENIERÍA INDUSTRIAL Y DE SISTEMAS',
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12,
         space_before=0, space_after=6)
add_para(doc, 'UNIDAD DE POSGRADO — MAESTRÍA EN INTELIGENCIA ARTIFICIAL',
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12,
         space_before=0, space_after=36)
add_para(doc, 'PROYECTO DE INVESTIGACIÓN',
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14,
         color=(0x1F, 0x38, 0x64), space_before=0, space_after=18)
add_para(doc,
    '«DETECCIÓN AUTOMATIZADA Y PRIORIZACIÓN DE EVENTOS ADVERSOS EN NOTAS CLÍNICAS '
    'MEDIANTE UN FLUJO DE PROCESAMIENTO (PIPELINE) DE NLP Y APRENDIZAJE AUTOMÁTICO — '
    'APLICACIÓN DE LA MATRIZ DE PRIORIZACIÓN DEL MODELO GEMSES SOBRE MIMIC-IV»',
    align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12,
    space_before=0, space_after=36)
add_para(doc, 'TRABAJO PARCIAL — CAPÍTULOS I Y II',
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12,
         space_before=0, space_after=24)
add_para(doc, 'Elaborado por:', align=WD_ALIGN_PARAGRAPH.CENTER,
         size=12, space_before=0, space_after=4)
add_para(doc, 'Carlos Pérez Pérez', align=WD_ALIGN_PARAGRAPH.CENTER,
         bold=True, size=12, space_before=0, space_after=18)
add_para(doc, 'Curso: MIA-303 — Proyectos de Investigación I',
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=0, space_after=4)
add_para(doc, 'Docente: Mg. María Tejada Begazo',
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=0, space_after=36)
add_para(doc, 'Lima — Perú, 2026',
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=0, space_after=0)

# Page break after portada
doc.add_page_break()

# ─── CAPÍTULO I ─────────────────────────────────────────────────────────────

add_heading(doc, 'CAPÍTULO I: PLANTEAMIENTO DEL PROBLEMA Y JUSTIFICACIÓN', level=0)

add_heading(doc, '1.1 Planteamiento del Problema', level=1)

# CP1: Párrafo expandido con datos RENIPRESS (26,663 IPRESS, 33M)
add_body(doc,
    'La seguridad del paciente constituye uno de los principales retos de los sistemas '
    'de salud contemporáneos. Según la Organización Mundial de la Salud (OMS, 2019), '
    'los eventos adversos derivados de la atención sanitaria figuran entre las diez '
    'principales causas de muerte y discapacidad en el mundo, generando además costos '
    'sustanciales para los sistemas asistenciales. En el Perú, el Registro Nacional de '
    'Instituciones Prestadoras de Servicios de Salud (RENIPRESS) registra 26,663 '
    'establecimientos activos que atienden a más de 33 millones de personas, distribuidos '
    'en doce tipos institucionales: sector privado (16,685 IPRESS), Gobierno Regional '
    '(8,380), MINSA (510), EsSalud (408), Sanidades de las Fuerzas Armadas y Policiales, '
    'SISOL, gobiernos locales y otros. EsSalud, que concentra la mayor densidad de '
    'hospitalizaciones de tercer nivel, aprobó en 2021 la Directiva GG-ESSALUD '
    '«Registro, Notificación y Gestión de los Eventos Relacionados con la Seguridad del '
    'Paciente (ERSP)», la cual incorpora explícitamente la Matriz de Priorización del '
    'Modelo de Gestión Moderna de los Servicios de Salud (GEMSES) como herramienta '
    'normativa para clasificar el nivel de riesgo y de gestión de los eventos.'
)

add_body(doc,
    'Sin embargo, el proceso de identificación y priorización depende todavía del '
    'registro voluntario de los profesionales de la salud y del análisis manual de las '
    'descripciones libres consignadas en las notas clínicas. Esta dependencia genera tres '
    'obstáculos persistentes: (i) la subnotificación voluntaria, asociada al temor de '
    'repercusiones laborales; (ii) la dependencia de descripciones en lenguaje natural '
    'que dificultan su codificación, agregación y análisis sistemático; y (iii) la falta '
    'de mecanismos automatizados que prioricen los eventos de mayor riesgo para '
    'escalarlos al nivel de gestión adecuado en tiempo oportuno.'
)

add_body(doc,
    'En el caso específico de EsSalud, la Directiva GG-ESSALUD-2021 establece un flujo '
    'formal de Registro → Notificación → Validación → Codificación → Priorización → '
    'Mitigación → Análisis de Causa Raíz → Plan de Acción. Aunque la directiva normaliza '
    'la Matriz de Priorización GEMSES como instrumento oficial, su aplicación efectiva '
    'requiere que personal de la Oficina de Gestión de la Calidad y Humanización procese '
    'manualmente cada notificación, lo que constituye un cuello de botella estructural '
    'que hace el proceso lento, costoso y vulnerable a la heterogeneidad en la '
    'interpretación.'
)

# CP2: Párrafo "demostrador técnico"
add_body(doc,
    'La presente investigación propone aprovechar los avances recientes en procesamiento '
    'de lenguaje natural (NLP) y aprendizaje automático para automatizar la detección de '
    '188 de los 231 eventos adversos del Anexo 02 de la Directiva GG-ESSALUD-2021 en '
    'notas clínicas no estructuradas, y la posterior aplicación de la Matriz de '
    'Priorización GEMSES. Dada la ausencia de acceso a registros institucionales '
    'peruanos, el estudio se desarrollará sobre el dataset abierto MIMIC-IV (Medical '
    'Information Mart for Intensive Care, versión IV). El estudio se concibe como un '
    'demostrador técnico de viabilidad: demuestra que la automatización GEMSES es '
    'realizable con datos clínicos reales, generando un protocolo documentado y '
    'replicable que las instituciones peruanas —cuyo sistema nacional aún carece de '
    'soporte normativo e infraestructura para esta automatización— podrán adoptar cuando '
    'se habiliten los accesos correspondientes.'
)

add_heading(doc, '1.2 Formulación del Problema', level=1)

# CP3: Formulación del problema con la cadena de 4 pasos de fallo
add_body(doc,
    'La gestión de eventos adversos en los servicios de salud presenta brechas críticas '
    'en la detección oportuna, la codificación consistente y la priorización basada en '
    'evidencia. Los modelos manuales dependen del criterio subjetivo del profesional y '
    'generan subnotificación sistemática. Más aún, incluso cuando un evento se reporta, '
    'la cadena de gestión presenta cuatro puntos de quiebre sucesivos: (i) el evento se '
    'registra pero no se analiza, por saturación del personal de calidad; (ii) se analiza '
    'pero no se prioriza, por ausencia de herramientas de clasificación estandarizada; '
    '(iii) se prioriza pero no se asigna un responsable con plazo definido; y (iv) se '
    'gestiona puntualmente pero no existe un indicador permanente de control de calidad '
    'que mida la recurrencia del evento. Este patrón impide que los sistemas de salud '
    'aprendan de sus errores y escalen la vigilancia frente al volumen de notas clínicas '
    'generadas en hospitales de tercer nivel.'
)

add_body(doc,
    'En consecuencia, el problema central se formula mediante la siguiente pregunta de '
    'investigación:'
)

add_body(doc,
    '¿En qué medida un pipeline integrado de procesamiento de lenguaje natural y '
    'aprendizaje automático supervisado, aplicado sobre notas clínicas no estructuradas '
    'del dataset MIMIC-IV, permite detectar eventos adversos y aplicar la Matriz de '
    'Priorización del Modelo GEMSES —incluyendo la asignación automática del Nivel de '
    'Gestión y del responsable institucional— con una concordancia significativamente '
    'mayor a la clasificación manual experta tradicional?',
    first_indent=False
)

add_heading(doc, '1.3 Justificación', level=1)

add_body(doc,
    'La presente investigación se justifica desde cuatro dimensiones complementarias:',
    first_indent=True
)

add_heading(doc, '1.3.1 Justificación Teórica', level=2)

add_body(doc,
    'El campo de la detección automática de eventos adversos en texto clínico ha '
    'avanzado sustancialmente en los últimos años, con trabajos basados en sistemas de '
    'supervisión débil (Ratner et al., 2017), modelos de transformers biomédicos (Lee '
    'et al., 2020; Alsentzer et al., 2019) y técnicas de negación clínica como NegEx '
    '(Chapman et al., 2001). Sin embargo, ningún trabajo previo ha articulado estos '
    'enfoques con la Matriz de Priorización del Modelo GEMSES ni con la taxonomía de '
    '231 eventos del Anexo 02 de la Directiva GG-ESSALUD-2021. Esta investigación cierra '
    'esa brecha teórica, aportando un marco de operacionalización que vincula las '
    'dimensiones de Tiempo, Calidad, Costos y Satisfacción del Modelo GEMSES con '
    'variables derivables de datos clínicos estructurados y no estructurados. Los pesos '
    'de impacto por dimensión ya están precalculados en el Anexo 02 para cada uno de los '
    '231 eventos; esta investigación los operacionaliza computacionalmente.'
)

add_body(doc,
    'Adicionalmente, el estudio contribuye al debate metodológico sobre la '
    'transferibilidad de modelos de NLP entrenados en inglés clínico hacia sistemas '
    'sanitarios de habla hispana, problema de creciente relevancia en la literatura de '
    'informática biomédica internacional.'
)

add_heading(doc, '1.3.2 Justificación Práctica', level=2)

add_body(doc,
    'Desde una perspectiva práctica, el sistema propuesto responde a una necesidad '
    'institucional concreta: el cuello de botella operativo que genera el procesamiento '
    'manual de notificaciones de eventos adversos en EsSalud. El pipeline automatizado '
    'permitiría procesar el volumen completo de notas clínicas de alta (331,793 notas en '
    'MIMIC-IV; equivalente a decenas de miles en EsSalud) sin incremento proporcional de '
    'recursos humanos especializados, reduciendo el tiempo entre la ocurrencia del evento '
    'y su priorización para la gestión correctiva.'
)

add_body(doc,
    'El modelo validado sobre MIMIC-IV generará un protocolo de implementación '
    'documentado y replicable, que podrá ser adaptado por la Oficina de Gestión de la '
    'Calidad y Humanización de EsSalud —o cualquier institución que opere bajo estándares '
    'GEMSES— cuando se habilite el acceso a datos clínicos institucionales para '
    'investigación. El impacto práctico esperado incluye: reducción de la '
    'subnotificación, mayor consistencia en la codificación, y escalamiento de la '
    'capacidad de vigilancia de la seguridad del paciente.'
)

add_heading(doc, '1.3.3 Justificación Metodológica', level=2)

add_body(doc,
    'El uso de MIMIC-IV como corpus de validación garantiza la reproducibilidad del '
    'estudio: es un dataset de acceso controlado pero abierto (PhysioNet credentialing), '
    'con 145,914 pacientes únicos, en formato estándar internacional, y ampliamente '
    'utilizado como benchmark en investigación de NLP clínico. La elección de este '
    'dataset —en lugar de datos institucionales peruanos no accesibles— no es una '
    'limitación sino una decisión metodológica deliberada que garantiza la trazabilidad, '
    'la posibilidad de revisión por pares y la comparabilidad con el estado del arte '
    'internacional.'
)

add_body(doc,
    'El diseño del pipeline combina supervisión débil para etiquetado a escala, '
    'detección de negación con NegEx, mapeo a códigos ICD-10, y comparación de múltiples '
    'clasificadores (TF-IDF + regresión logística, LinearSVC, modelos basados en '
    'transformers). La validación mediante panel experto (n = 70 notas) con métricas de '
    'concordancia (kappa de Cohen) proporciona evidencia estadística rigurosa sobre la '
    'capacidad del sistema, con un IC95% Wilson calculado correctamente para muestras '
    'pequeñas.'
)

add_heading(doc, '1.3.4 Justificación Social', level=2)

add_body(doc,
    'La OMS estima que en países de ingresos medios y bajos, la carga de eventos '
    'adversos evitables equivale a 134 millones de episodios anuales con 2.6 millones '
    'de muertes asociadas (OMS, 2019). Perú, como país de ingreso medio-alto con un '
    'sistema de seguridad social en expansión, enfrenta este problema de manera aguda: '
    'la red asistencial de EsSalud atiende a más de 13 millones de asegurados con '
    'capacidad limitada de vigilancia activa de seguridad.'
)

add_body(doc,
    'Un sistema automatizado de detección y priorización contribuye directamente a la '
    'protección de pacientes vulnerables, a la rendición de cuentas institucional y al '
    'aprendizaje organizacional. La generación de alertas tempranas sobre eventos de '
    'nivel Rojo (alta prioridad GEMSES) permitiría activar protocolos de mitigación '
    'antes de que el daño escale, con un impacto social concreto y medible en términos '
    'de morbimortalidad prevenible.'
)

# ─── CAPÍTULO II ─────────────────────────────────────────────────────────────

doc.add_page_break()

add_heading(doc, 'CAPÍTULO II: OBJETIVOS', level=0)

add_heading(doc, '2.1 Objetivo General', level=1)

add_body(doc,
    'Diseñar, implementar y validar un sistema basado en procesamiento de lenguaje '
    'natural y aprendizaje automático que detecte eventos adversos a partir de notas '
    'clínicas no estructuradas del dataset MIMIC-IV y los priorice automáticamente '
    'aplicando la Matriz de Priorización del Modelo GEMSES —incluyendo el cálculo '
    'automatizado del Nivel de Gestión y la asignación del responsable institucional— '
    'con miras a su transferencia metodológica a sistemas sanitarios de habla hispana y, '
    'en particular, a la red asistencial de EsSalud.'
)

add_heading(doc, '2.2 Objetivos Específicos', level=1)

add_body(doc,
    'Para alcanzar el objetivo general se plantean los siguientes objetivos específicos:',
    first_indent=True
)

# OE1 — updated to 188 de 231 events
add_objective(doc,
    'OE1 — Construir un corpus etiquetado de eventos adversos a partir de notas clínicas '
    'no estructuradas de MIMIC-IV, utilizando como marco de etiquetado los 188 eventos '
    'detectables (de los 231 del Anexo 02 de la Directiva GG-ESSALUD-2021), clasificados '
    'por cuatro niveles de detección: (A) ICD-10 + NLP, (B) patrones de texto libre, '
    '(C) variables estructuradas calculadas y (D) auditoría de completitud de historia '
    'clínica.'
)

# OE2 — unchanged
add_objective(doc,
    'OE2 — Implementar y comparar al menos tres modelos de procesamiento de lenguaje '
    'natural para la detección y clasificación de eventos adversos en texto libre, '
    'incluyendo modelos basados en transformers biomédicos (BioBERT, ClinicalBERT) y un '
    'modelo base clásico de referencia (TF-IDF + regresión logística).'
)

# OE3 — CP4: rewritten to automate GEMSES formula
add_objective(doc,
    'OE3 — Automatizar la priorización GEMSES aplicando, para cada evento detectado, la '
    'fórmula Nivel de Riesgo = Frecuencia × Impacto usando los pesos de impacto ya '
    'establecidos en el Anexo 02 para cada uno de los 231 eventos (dimensiones: Estancia '
    'Hospitalaria, Complicaciones, Costos, Satisfacción del Paciente e Impacto '
    'Ambiental), y determinar automáticamente el Nivel de Gestión (Verde / Amarillo / '
    'Rojo) y el responsable institucional correspondiente.'
)

# OE4 — CP5: simplified language
add_objective(doc,
    'OE4 — Validar la precisión del pipeline comparando sus resultados contra la '
    'revisión manual de un panel de expertos sobre una muestra de 70 notas clínicas '
    'seleccionadas de forma estratificada, utilizando métricas estándar de concordancia '
    'e índices de clasificación para determinar si el sistema automatizado alcanza un '
    'nivel de acuerdo aceptable con el juicio experto.'
)

# OE5 — unchanged
add_objective(doc,
    'OE5 — Documentar el método de transferencia del pipeline a sistemas sanitarios de '
    'habla hispana, identificando los retos lingüísticos, taxonómicos y de localización '
    'que deberán abordarse para su implementación en EsSalud u otras instituciones '
    'equivalentes.'
)

# CP6: Section 2.3 ELIMINADA — no section header, articulación woven in
add_body(doc,
    'Los objetivos específicos OE1 y OE2 responden al obstáculo de la dependencia del '
    'análisis manual mediante la construcción de corpus etiquetado y la comparación de '
    'modelos. OE3 operacionaliza el instrumento GEMSES con sus pesos de impacto '
    'oficiales, automatizando la cadena completa Frecuencia × Impacto → Nivel de '
    'Gestión → Responsable. OE4 provee la evidencia empírica de validez con rigor '
    'estadístico. OE5 conecta el resultado de laboratorio con el impacto social y '
    'práctico, cerrando el ciclo de investigación aplicada y garantizando la coherencia '
    'global del proyecto.',
    first_indent=True
)

# ─── Save ─────────────────────────────────────────────────────────────────────

out_path = 'C:/MIMIC/tesis/01_plan_de_tesis/Trabajo_Parcial_Cap1_Cap2_v2.docx'
doc.save(out_path)
print(f'OK: {out_path}')

# Quick content verification
doc2 = Document(out_path)
paras = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
print(f'Paragraphs: {len(paras)}')
checks = [
    ('CP1-RENIPRESS', '26,663'),
    ('CP1-tipos', 'doce tipos institucionales'),
    ('CP2-demostrador', 'demostrador técnico de viabilidad'),
    ('CP3-cadena4', 'cuatro puntos de quiebre'),
    ('CP3-noasigna', 'no se asigna un responsable'),
    ('CP3-indicador', 'indicador permanente de control de calidad'),
    ('CP4-OE3-formula', 'Frecuencia × Impacto'),
    ('CP4-OE3-pesos', 'pesos de impacto ya establecidos en el Anexo 02'),
    ('CP5-OE4-simple', 'revisión manual de un panel de expertos'),
    ('CP6-sin2.3', '2.3'),
    ('OE1-188', '188 eventos detectables'),
    ('231total', '231 eventos'),
]

full_text = '\n'.join(paras)
print('\n--- Verificación CP1–CP6 ---')
for label, term in checks:
    found = '✓' if term in full_text else '✗ FALTA'
    if label == 'CP6-sin2.3':
        found = '✓ eliminada' if '2.3' not in full_text else '✗ TODAVÍA PRESENTE'
    print(f'  {label}: {found}')
